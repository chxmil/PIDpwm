"""
Builds Research/Material_Classification_Methods_2026-05-13.docx from the markdown
report of the same name, embedding figures from Research/figures/.

Pipeline: markdown -> HTML (python `markdown` lib) -> python-docx via BeautifulSoup.

Figure placement: each heading is matched against FIGURE_INSERTS by exact text;
matching figures are inserted (with caption) immediately after the heading's
introductory paragraph.

Run:
    python "Code Store/build_methods_docx.py"
"""
import os
import re
import sys
from html import unescape

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE     = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH  = os.path.join(BASE, "Research", "Material_Classification_Methods_2026-05-13.md")
FIG_DIR  = os.path.join(BASE, "Research", "figures")
OUT_PATH = os.path.join(BASE, "Research", "Material_Classification_Methods_2026-05-13.docx")

# (figure_file, caption) — inserted after the FIRST paragraph of the named heading.
FIGURE_INSERTS = {
    "1. Data Acquisition": [
        ("fig1_example_signal.png",
         "Figure 1. Example signal trace from a Hard-material grip. Top to bottom: "
         "raw resistance, shifted conductance, joint angle, predicted force. "
         "Red line marks first contact (is_press=1); red-shaded region is the PID grip phase."),
    ],
    "4. Feature Engineering": [
        ("fig2_class_mean_trajectories.png",
         "Figure 2. Mean post-contact trajectories per material class. Left: shifted "
         "conductance over the first 2 seconds after contact. Right: change in joint "
         "angle relative to contact. Hard, Medium, Soft are visually separable in both "
         "channels but the gap narrows under PID overshoot."),
    ],
    "4.1 Phase A — RF, Hand-Crafted Scalars": [
        ("fig5_rf_feature_importance.png",
         "Figure 3. Random-Forest v4 feature importance (impurity decrease). "
         "f_peak and stiffness_proxy dominate; delta_pos_max is effectively dead "
         "and a candidate for removal."),
    ],
    "5. Window Extraction (Time-Series Classifiers)": [
        ("fig6_example_windows.png",
         "Figure 4. Example (40, 5) windows for one Hard, one Medium, and one Soft "
         "trial after per-channel z-score normalisation. Time axis runs left-to-right "
         "in 50 ms bins. Visual differences between classes drive the 1D-CNN."),
    ],
    "11. Metrics": [
        ("fig3_confusion_matrices.png",
         "Figure 5. Confusion matrices on the full data_logs/datasets/ corpus (in-sample). "
         "RF v4 left, CNN-PID v2 right."),
        ("fig4_per_class_f1.png",
         "Figure 6. Per-class F1 score comparison — RF v4 vs CNN-PID v2 on the same "
         "back-test set."),
    ],
}


# --- python-docx helpers -----------------------------------------------------

def _set_font(run, name="Calibri", size_pt=11, bold=False, italic=False, color=None):
    run.font.name = name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_paragraph_with_runs(doc, html_node, style=None):
    """Append a paragraph whose runs come from inline children of html_node.

    Handles <strong>, <em>, <code>, <a>, and bare text.
    """
    p = doc.add_paragraph(style=style)
    _emit_inline(p, html_node)
    return p


def _emit_inline(p, html_node, bold=False, italic=False, code=False):
    for child in html_node.children:
        name = getattr(child, "name", None)
        if name is None:
            text = unescape(str(child))
            if not text:
                continue
            r = p.add_run(text)
            if code:
                _set_font(r, name="Consolas", size_pt=10, bold=bold, italic=italic,
                          color=(0xb2, 0x4a, 0x1f))
            else:
                _set_font(r, bold=bold, italic=italic)
        elif name in ("strong", "b"):
            _emit_inline(p, child, bold=True,  italic=italic, code=code)
        elif name in ("em", "i"):
            _emit_inline(p, child, bold=bold, italic=True, code=code)
        elif name == "code":
            _emit_inline(p, child, bold=bold, italic=italic, code=True)
        elif name == "a":
            _emit_inline(p, child, bold=bold, italic=italic, code=code)
        else:
            _emit_inline(p, child, bold=bold, italic=italic, code=code)


def _add_figure(doc, image_path, caption_text, width_inches=6.2):
    if not os.path.exists(image_path):
        print(f"  [warn] missing figure: {image_path}")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(image_path, width=Inches(width_inches))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run(caption_text)
    _set_font(r, size_pt=10, italic=True, color=(0x44, 0x44, 0x44))


def _add_code_block(doc, code_text):
    for line in code_text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line if line else " ")
        _set_font(r, name="Consolas", size_pt=10, color=(0x20, 0x20, 0x40))


def _add_table(doc, html_table):
    headers = [th.get_text(strip=True) for th in html_table.find_all("th")]
    rows    = []
    for tr in html_table.find_all("tr"):
        cells = tr.find_all("td")
        if cells:
            rows.append([td.get_text(" ", strip=True) for td in cells])
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    if headers:
        for j, h in enumerate(headers[:n_cols]):
            cell = table.rows[0].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(h)
            _set_font(r, bold=True, size_pt=10)
    for i, row in enumerate(rows, start=1):
        for j in range(n_cols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(row[j] if j < len(row) else "")
            _set_font(r, size_pt=10)


def _heading_key(heading_text):
    """Normalise a heading text for matching against FIGURE_INSERTS keys."""
    return heading_text.strip()


# --- main conversion ---------------------------------------------------------

def build():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()
    html = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "sane_lists"],
    )
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    # Default style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Material Classification — Full Method Report")
    _set_font(r, size_pt=18, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("AI-Adaptive PID Control for Tactile Robotic Grippers · 2026-05-13")
    _set_font(r, size_pt=11, italic=True, color=(0x55, 0x55, 0x55))
    doc.add_paragraph()

    pending_inserts = []   # [(file, caption), ...] to insert after the next paragraph

    for node in soup.children:
        name = getattr(node, "name", None)
        if name is None:
            text = str(node).strip()
            if text:
                doc.add_paragraph(text)
            continue

        if name == "h1":
            doc.add_heading(node.get_text(strip=True), level=1)
        elif name == "h2":
            txt = node.get_text(strip=True)
            doc.add_heading(txt, level=2)
            key = _heading_key(txt)
            if key in FIGURE_INSERTS:
                pending_inserts.extend(FIGURE_INSERTS[key])
        elif name == "h3":
            txt = node.get_text(strip=True)
            doc.add_heading(txt, level=3)
            key = _heading_key(txt)
            if key in FIGURE_INSERTS:
                pending_inserts.extend(FIGURE_INSERTS[key])
        elif name == "h4":
            doc.add_heading(node.get_text(strip=True), level=4)
        elif name == "p":
            _add_paragraph_with_runs(doc, node)
            # Insert any pending figures after this paragraph
            for fn, cap in pending_inserts:
                _add_figure(doc, os.path.join(FIG_DIR, fn), cap)
            pending_inserts.clear()
        elif name == "ul":
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                _emit_inline(p, li)
        elif name == "ol":
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Number")
                _emit_inline(p, li)
        elif name == "pre":
            code_text = node.get_text()
            _add_code_block(doc, code_text)
        elif name == "table":
            _add_table(doc, node)
        elif name == "hr":
            doc.add_paragraph()  # blank line as separator
        elif name == "blockquote":
            for child in node.find_all(["p"], recursive=False):
                p = doc.add_paragraph(style="Quote")
                _emit_inline(p, child)
        else:
            doc.add_paragraph(node.get_text(" ", strip=True))

    # Flush any unconsumed inserts
    for fn, cap in pending_inserts:
        _add_figure(doc, os.path.join(FIG_DIR, fn), cap)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
    print(f"Size : {os.path.getsize(OUT_PATH)//1024} KB")


if __name__ == "__main__":
    sys.exit(build())
